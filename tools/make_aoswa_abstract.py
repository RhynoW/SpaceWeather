"""tools.make_aoswa_abstract — 產生 AOSWA 2026 摘要 docx。

格式取自 docs/AOSWA2026_Abstract_Template.docx：A4、單頁、300–800 字，
Times New Roman，標題 14pt 粗體置中、作者 12pt 置中、單位 10pt 斜體置中、
內文 12pt 左右對齊、5 個關鍵詞。

**範本明訂標題與摘要不得使用符號、特殊字元、註腳或數學式**，
故內文一律以文字敘述取代希臘字母、箭頭與運算符號；本檔以純 ASCII 撰寫，
存檔後可用 `python -c` 掃描非 ASCII 字元複驗。

作者姓名與單位為佔位符，投稿前須填寫。掛名須經共同作者同意。
"""

import re
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH as A
from docx.oxml.ns import qn

TITLE = ("An Open-Source Multi-Domain Space Weather Visualization Application: "
         "Observed Thermospheric Density Criteria from FORMOSAT-7 "
         "and a Silent Dependency Failure")

AUTHORS = "A. A. Author1*, B. B. Second-Author2"
AFFIL = ["1Affiliation to be completed, City, Taiwan",
         "2Affiliation to be completed, City, Taiwan"]
EMAIL = "*Corresponding author's email: rhynowu@gmail.com"

BODY = [
 "Space weather services in the Asia-Oceania region need tools that turn heterogeneous "
 "observations into judgements a non-specialist can act on. We report an open-source web "
 "application presenting space weather risk for three impact domains: high-frequency radio "
 "communication, GNSS positioning, and low-Earth-orbit prediction. It is built on a "
 "bitemporal store in which every record carries both a physical valid time and the time it "
 "was ingested, so past events replay exactly as they appeared to an operator at the time.",

 "Two design decisions proved more consequential than the visualization itself. The system "
 "distinguishes three states rather than two: a criterion may be evaluated and found below "
 "threshold, evaluated only in part because some inputs are missing, or not evaluable at all. "
 "Merging the latter two into a green indicator would present absence of data as absence of "
 "hazard. Every derived quantity also carries a provenance label recording whether it is "
 "observed, modelled, inferred from a proxy, or unavailable.",

 "Integrating two FORMOSAT-7 products moved two domains off proxy indicators. Onboard "
 "scintillation data gave the first observed criterion for GNSS positioning, with amplitude "
 "scintillation index values up to 1.17 over the Taiwan sector. Precise orbit determination "
 "gave the first observed criterion for orbit prediction, whose four rules had previously all "
 "been thresholds on geomagnetic indices. We derive an orbit-averaged drag decay rate from "
 "the semi-major axis and divide it by the value expected at the same solar radio flux under "
 "quiet conditions. The ballistic coefficient cancels in that ratio, so no spacecraft "
 "cross-sectional area, drag coefficient, or dry mass is needed, and none of these is public "
 "for this satellite. Two conditions proved essential: the averaging window must span an "
 "integer number of orbital periods, because the short-period variation of the osculating "
 "semi-major axis is far larger than the daily decay and otherwise aliases into an apparent "
 "trend; and the baseline must remove the solar flux contribution, because over our test "
 "window the solar radio flux rose from 132 to 238 and a rolling baseline attributes that "
 "rise to geomagnetic activity.",

 "Two results follow. During the May 2024 Gannon storm the observed enhancement peaked at "
 "4.0, while an empirical model evaluated along the same track gave 2.3; the observed to "
 "modelled ratio rises monotonically with enhancement, from 0.85 in quiet conditions to 1.7 "
 "at the largest values, indicating a compressed model storm response. Separately, the "
 "density response is strongly nonlinear in geomagnetic activity: for planetary K index "
 "between 6 and 7 the median enhancement is only 1.19, reaching 2.34 only above 7. Over the "
 "same window the index-based rules fired fifteen times and the observed rules four times, "
 "all during the storm, so the proxy over-alerts at moderate disturbance.",

 "A finding of wider relevance emerged during this work. An installed astronomy library was "
 "incompatible with the installed numerical array package, and its time and coordinate "
 "modules could not be imported at all. Two independent causes were present: a removed array "
 "function, and a private symbol imported across a package boundary and since renamed. The "
 "common workaround addresses only the first. More seriously, a calling routine caught the "
 "import failure and re-raised it advising the user to install the library, which was already "
 "installed, so the diagnostic pointed away from the true cause. The affected pipeline could "
 "not run and its results table stayed empty with no visible error. After upgrading the "
 "library we implemented an independent coordinate transformation and pinned it against the "
 "library across 2636 orbit points, agreeing to 0.08 m on average. We argue that dependency "
 "verification belongs inside the data quality regime of operational space weather software, "
 "because this failure class yields no wrong number, only a missing one.",

 "The application includes a multilingual educational module for students aged twelve to "
 "eighteen in Traditional Chinese, Japanese, English and Malay. Source code, configuration "
 "and the data source inventory are public.",
]

KEYWORDS = ("Keywords: space weather visualization, situational awareness, data provenance, "
            "thermospheric density, software reproducibility")

d = docx.Document()
s = d.sections[0]
s.page_width, s.page_height = Inches(8.27), Inches(11.68)
s.left_margin, s.right_margin = Inches(0.93), Inches(0.92)
s.top_margin, s.bottom_margin = Inches(1.10), Inches(0.19)

st = d.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(12)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
st.paragraph_format.space_after = Pt(0)

def _run(p, text, size, bold, italic, sup=False):
    r = p.add_run(text); r.font.name = "Times New Roman"
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if sup:
        r.font.superscript = True
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def para(text, size=12, bold=False, italic=False, align=A.JUSTIFY, after=0, sup=None):
    """sup 為正規式；符合的片段以上標排版（作者與單位的編號）。"""
    p = d.add_paragraph(); p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if sup is None:
        _run(p, text, size, bold, italic)
        return p
    for part in re.split("(" + sup + ")", text):
        if part:
            # 切出來的純數字片段即為編號記號
            _run(p, part, size, bold, italic, sup=part.isdigit())
    return p

para(TITLE, size=14, bold=True, align=A.CENTER, after=12)
para(AUTHORS, size=12, align=A.CENTER, after=6, sup=r"(?<=r)1(?=\*)|(?<=r)2$")
for a in AFFIL:
    para(a, size=10, italic=True, align=A.CENTER, sup=r"^[12]")
para(EMAIL, size=10, align=A.CENTER, after=12)
for b in BODY:
    para(b, after=8)
para("", after=4)
para(KEYWORDS)

out = "docs/AOSWA2026_Abstract_SpaceWeather_App.docx"
d.save(out)
print(out, "words in body =", sum(len(b.split()) for b in BODY))
