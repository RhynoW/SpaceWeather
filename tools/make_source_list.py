"""tools/make_source_list.py — 由設定產生資料來源清單 docx。

為什麼要程式產生而不是手工維護：來源、影像與動畫都定義在 `configs/`，
手工抄一份到 Word 就會在下次新增來源時悄悄過期——而且沒有任何徵兆。
本工具一律從 `catalog()`／`imagery()`／`animations()` 取當下的內容，
產生的清單不可能與系統不一致。

**標註是這份文件的重點**：本系統整合的全是外部機構產製的資料，
引用時有標註義務。每張表都帶產製者與使用條款欄位。

用法：
    python tools/make_source_list.py
    python tools/make_source_list.py --out docs/SpaceWeather資料來源清單20260818.docx
"""

from __future__ import annotations

import argparse
import sys
from datetime import datetime, timezone
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "packages"))
sys.path.insert(0, str(ROOT))

GITHUB = "https://github.com/RhynoW/SpaceWeather"
DEMO = "https://spaceweather.streamlit.app/"

# 每個來源「為什麼需要它」——這是設定檔裡沒有、也不該有的判斷，
# 屬於文件層的敘述，因此留在這裡以 source_id 對應。
WHY = {
    "celestrak_sw_all": "STK/GMAT 驅動檔的資料本體。CSSI 格式的權威來源，2021→2041 含月預測",
    "gfz_nowcast": "Kp 的原始產製者（CelesTrak 亦轉載自此）。備援兼近即時",
    "swpc_xray": "閃焰的即時強度，R 級判定依據。1 分鐘節奏",
    "swpc_solarwind_mag": "整條因果鏈最關鍵的參數。南向 Bz 才會有地磁暴",
    "swpc_solarwind_plasma": "與 Bz 併看才有意義。DSCOVR/ACE 在 L1，約 30–60 分鐘提前量",
    "swpc_kp_estimated": "1 分鐘更新，比正式 Kp（3 小時）早知道擾動起來了",
    "swpc_protons": "S 級判定。極區 HF 中斷與衛星單粒子翻轉風險",
    "swpc_xray_flares": "事件化的閃焰（相對於連續通量）",
    "swpc_solar_regions": "系統中唯一對閃焰有提前量的資訊，但僅產生 L1 提示",
    "swpc_ovation": "擾動深入到什麼緯度的直觀指標",
    "swpc_geomag_forecast": "預報引擎的對照基準——本系統模型必須與它比較",
    "swpc_27day_outlook": "27 日太陽自轉週期的復現預報",
    "kyoto_dst": "環電流強度。比 Kp 更直接反映注入磁層的能量",
    "cwa_swoo": "唯一的在地實測。使 GNSS_PNT 網域從規則全部 unavailable 變成可判定",
    "swpc_drap": "HF 中斷的直接判據。架構書原列「需協調」，實為公開產品",
    "gfz_hp30": "把暴起始時刻的解析度從 3 小時提升到 30 分鐘",
    "omni2_hourly": "預報引擎唯一可用的長期訓練資料（1963 年起）",
    "tw_gnss_tec": "在地電離層閃爍實測",
    "tw_magnetometer": "在地地磁擾動實測",
    "tw_ionosonde": "在地電離層探測",
}

EXCLUDE_WHY = {
    "gfz_hp30": "單獨就要約 46 秒，佔全部擷取時間六成。改由手動「完整」按鈕或排程處理",
    "omni2_hourly": "六年份歷史回填，且本身是事後重整資料（發布延遲數週至數月），無即時價值",
}


def _shade_header(row) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for cell in row.cells:
        tcpr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "1B4F79")
        tcpr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = None
                from docx.shared import RGBColor

                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _table(doc, headers: list[str], rows: list[list[str]], widths=None):
    from docx.shared import Cm, Pt

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for cell, head in zip(t.rows[0].cells, headers):
        cell.text = head
    _shade_header(t.rows[0])
    for r in rows:
        cells = t.add_row().cells
        for cell, value in zip(cells, r):
            cell.text = str(value)
    for row in t.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for cell, w in zip(row.cells, widths):
                cell.width = Cm(w)
    return t


def build(out_path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt

    from swx_core import animations, catalog, imagery
    from services.ingest.refresh import HEAVY_SOURCES, live_sources

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft JhengHei"
    style.font.size = Pt(10)

    today = datetime.now(timezone.utc).astimezone().strftime("%Y/%m/%d")
    h = doc.add_heading(f"SpaceWeather 資料來源清單（{today}）", level=0)
    h.runs[0].font.name = "Microsoft JhengHei"

    doc.add_paragraph(f"GitHub　{GITHUB}")
    doc.add_paragraph(f"Live Demo　{DEMO}")

    srcs = list(catalog())
    ready = [s for s in srcs if s.status == "ready"]
    planned = [s for s in srcs if s.status != "ready"]
    auto = set(live_sources())
    manual = [s for s in ready if s.source_id not in auto]

    p = doc.add_paragraph()
    p.add_run(
        f"{len(ready)} 個 ready（{len(auto)} 個進自動更新、{len(manual)} 個手動）"
        f"＋ {len(planned)} 個 planned。"
    ).bold = True

    note = doc.add_paragraph()
    note.add_run("引用義務：").bold = True
    note.add_run(
        "本系統整合的全部為外部機構產製的資料。引用本系統的任何數字時，"
        "須一併標註原始產製者。下表的「使用條款」欄取自 configs/sources.yaml 的 "
        "attribution 區段，由契約測試確保不得遺漏。"
    )

    # ── 自動更新 ──
    doc.add_heading(f"一、自動更新的 {len(auto)} 個來源", level=1)
    rows = []
    for i, s in enumerate([x for x in ready if x.source_id in auto], 1):
        a = s.raw.get("attribution", {})
        rows.append([i, s.source_id, a.get("provider", ""),
                     ", ".join(s.provides), WHY.get(s.source_id, ""),
                     a.get("terms", "")])
    _table(doc, ["#", "來源", "產製者", "提供參數", "為什麼需要它", "使用條款"],
           rows, widths=[0.8, 3.0, 3.4, 3.6, 5.2, 4.8])

    # ── 手動 ──
    doc.add_heading(f"二、不進自動更新的 {len(manual)} 個", level=1)
    rows = []
    for s in manual:
        a = s.raw.get("attribution", {})
        rows.append([s.source_id, a.get("provider", ""), ", ".join(s.provides),
                     EXCLUDE_WHY.get(s.source_id, ""), a.get("terms", "")])
    _table(doc, ["來源", "產製者", "提供參數", "為什麼排除", "使用條款"],
           rows, widths=[3.0, 3.4, 3.6, 5.4, 4.4])

    # ── 待協調 ──
    doc.add_heading(f"三、待協調的 {len(planned)} 個", level=1)
    rows = [[s.source_id, (s.raw.get("attribution") or {}).get("provider", ""),
             WHY.get(s.source_id, ""), (s.raw.get("attribution") or {}).get("terms", "")]
            for s in planned]
    _table(doc, ["來源", "對口", "內容", "取得方式"], rows, widths=[3.4, 4.0, 5.4, 4.0])

    # ── 影像 ──
    imgs = imagery()
    doc.add_heading(f"四、影像來源（{len(imgs)} 張）", level=1)
    doc.add_paragraph(
        "一律直接連結產製者網址、不下載轉存——確保呈現的是對方當下的版本，"
        "也避免衍生重製散布的授權問題。標「模式輸出」者非觀測資料。"
    )
    rows = [[i["title"], i.get("instrument", ""), i["attribution"].get("provider", ""),
             i["attribution"].get("terms", "")] for i in imgs]
    _table(doc, ["影像", "儀器／模式", "產製者", "使用條款"], rows,
           widths=[4.0, 3.8, 4.6, 4.4])

    # ── 動畫 ──
    anims = animations()
    videos = [a for a in anims if a["kind"] == "video"]
    frames = [a for a in anims if a["kind"] == "frames"]
    doc.add_heading(f"五、動畫來源（{len(anims)} 段：{len(videos)} 支 MP4、"
                    f"{len(frames)} 段逐幀）", level=1)
    doc.add_paragraph(
        "作法依單幀大小決定：高解析度序列採產製者預先編碼的 MP4"
        "（SUVI 304Å 單幀 1.1 MB × 359 幀 = 397 MB，逐幀載入不切實際）；"
        "單幀小於 100 KB 者由 SWPC 幀索引即時組成，永遠是最新的一段。"
    )
    rows = []
    for a in anims:
        size = f"{a['approx_mb']} MB" if a.get("approx_mb") else f"{a.get('max_frames', '?')} 幀"
        rows.append([a["title"], a.get("instrument", ""),
                     "MP4" if a["kind"] == "video" else "逐幀", size,
                     a["attribution"].get("provider", "")])
    _table(doc, ["動畫", "儀器／模式", "型式", "量體", "產製者"], rows,
           widths=[4.2, 3.8, 1.6, 1.8, 5.4])

    # ── 授權特別說明 ──
    doc.add_heading("六、授權特別說明", level=1)
    warn = doc.add_paragraph()
    warn.add_run("中央氣象署 SWOO（cwa_swoo）").bold = True
    warn.add_run(
        "：該端點為 SWOO 網站前端自用，非公開發布之 API。本案經成功大學合作管道"
        "取得中央氣象署授權後使用。第三方不得逕行取用，須另行取得授權。"
        "引用時須標註中央氣象署為資料產製者。原始 JSON 不進版控。"
    )
    doc.add_paragraph(
        "其餘來源：NOAA/SWPC 與 NASA 為美國政府公眾領域資料；"
        "GFZ 為 CC BY 4.0（須引用對應論文）；WDC Kyoto 學術使用免費但須依其引用規範"
        "並註明所用為暫定值；CelesTrak 免費公開但請標註。"
    )

    doc.add_paragraph()
    tail = doc.add_paragraph()
    tail.add_run(
        "本清單由 tools/make_source_list.py 依 configs/ 的實際設定產生，"
        "不手工維護；新增來源後重新執行即可同步。"
    ).italic = True

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="產生資料來源清單 docx")
    ap.add_argument("--out", default=None)
    args = ap.parse_args(argv)

    default = ROOT / "docs" / f"SpaceWeather資料來源清單{datetime.now():%Y%m%d}.docx"
    out = Path(args.out) if args.out else default
    path = build(out)
    print(f"已產生 {path}（{path.stat().st_size / 1024:.0f} KB）")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
