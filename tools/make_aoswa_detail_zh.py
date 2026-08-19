"""tools.make_aoswa_detail_zh — 產生 AOSWA 2026 投稿的中文詳細版（內部討論用）。

與 `make_aoswa_abstract.py` 產出的投稿摘要是**兩份不同文件**：

  投稿摘要   英文、單頁、300–800 字、不得使用符號與數學式
  本文件     中文、不限長度、含方法細節、實測數字、已修正的錯誤與待討論事項

**為何要含「已修正的錯誤」**：同儕要判斷結論可不可信，需要知道哪些地方
查過、哪些地方曾經算錯又怎麼發現的。只列成功的結果無法支持審閱。

以腳本產生而非直接編輯二進位檔，數字更新時可重跑並以 git diff 檢視。
"""

from __future__ import annotations

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH as A
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

LATIN = "Times New Roman"
CJK = "標楷體"

TITLE = "太空天氣視覺化系統：由福衛七號精密定軌建立實測熱氣層密度判據"
SUBTITLE = "AOSWA 2026 投稿之中文詳細版　—　內部討論用，非投稿文本"
META = "版本 2026-08-19　·　撰寫：SpaceWeather 專案　·　狀態：待同儕審閱"


def _run(p, text, *, size=12, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.name = LATIN
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r._element.rPr.rFonts.set(qn("w:eastAsia"), CJK)
    return r


def _emit(p, text, *, size, bold, italic):
    """支援 **粗體** 標記。不解析的話星號會照字面印在文件上。"""
    for i, part in enumerate(text.split("**")):
        if part:
            _run(p, part, size=size, bold=bold or (i % 2 == 1), italic=italic)


def para(d, text="", *, size=12, bold=False, italic=False,
         align=A.JUSTIFY, after=6, indent=0.0):
    p = d.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if text:
        _emit(p, text, size=size, bold=bold, italic=italic)
    return p


def heading(d, text, level=1):
    sizes = {1: 15, 2: 13, 3: 12}
    p = para(d, text, size=sizes[level], bold=True, align=A.LEFT,
             after=6 if level > 1 else 8)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    return p


def bullets(d, items, *, indent=0.25):
    for it in items:
        p = para(d, "・" + it, after=3, indent=indent)
    return p


def table(d, header, rows, *, widths=None):
    t = d.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.autofit = True
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ""
        _run(cell.paragraphs[0], h, size=10.5, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            _run(cells[i].paragraphs[0], str(v), size=10.5)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    para(d, "", after=4)
    return t


def build() -> docx.Document:
    d = docx.Document()
    s = d.sections[0]
    s.page_width, s.page_height = Inches(8.27), Inches(11.69)
    for m in ("left_margin", "right_margin"):
        setattr(s, m, Inches(1.0))
    s.top_margin, s.bottom_margin = Inches(1.0), Inches(1.0)

    st = d.styles["Normal"]
    st.font.name = LATIN
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), CJK)

    para(d, TITLE, size=17, bold=True, align=A.CENTER, after=4)
    para(d, SUBTITLE, size=11, align=A.CENTER, after=2)
    para(d, META, size=10, italic=True, align=A.CENTER, after=12)

    # ── 一、這份文件要討論什麼 ──────────────────────────────────
    heading(d, "一、這份文件要討論什麼")
    para(d,
         "本專案的軌道預報網域原本四條分級規則全部以 Kp／Ap 為門檻。那是地磁"
         "代理指標——地磁擾動強不等於熱氣層密度一定高，只是相關。本次工作以福衛"
         "七號精密定軌反演出實測的密度增強倍數，使該網域首次具備非代理的判定"
         "基礎，並據以檢驗既有代理判據的偏差。")
    para(d,
         "請同儕特別針對第五節（方法）與第八節（待確認事項）給意見。"
         "第七節列出本次工作中曾經算錯、後來如何發現與修正的三個地方——"
         "列出它們是為了讓審閱者能判斷結論的可信程度，不是為了自陳。", after=8)

    # ── 二、系統概述 ────────────────────────────────────────────
    heading(d, "二、系統概述")
    para(d,
         "系統是一個開源網頁應用，將太空天氣風險呈現為三個衝擊網域：短波通信"
         "（HF_COMM）、衛星定位（GNSS_PNT）、軌道預報（ORBIT_PREDICTION）。"
         "資料層採雙時間軸儲存，每筆記錄同時帶物理有效時刻（valid_time）與"
         "入庫時刻（ingest_time），因此歷史事件可以「當時已知的資料」重播，"
         "不會有前視偏差。")
    para(d, "兩個設計決定的影響大於視覺化本身：", after=4)
    bullets(d, [
        "**三態而非兩態**：判據可能「已評估且未達門檻」「僅部分可評估」"
        "「完全無法評估」。把後兩者併入綠燈，等於把「沒有資料」呈現為「沒有危害」。",
        "**證據來源標記**：每個衍生量都帶 observed／modelled／proxy／unavailable，"
        "使用者看到的不只是等級，還有該等級的證據基礎。網域層級取其中最弱的一項。",
    ])

    heading(d, "三網域的判據現況", 2)
    table(d,
          ["網域", "實測判據", "代理判據", "現況"],
          [["HF_COMM", "GOES X 射線、質子通量", "—", "判據齊備"],
           ["GNSS_PNT", "福衛七號掩星 S4、TEC", "Kp、閃焰機率",
            "L3 需 S4＋ROTI；只有 S4 時回報 partial"],
           ["ORBIT_PREDICTION", "DRAG_ENHANCEMENT（本次新增）", "Kp、Ap",
            "兩組並用：Kp 給前導，實測給量值"]],
          widths=[1.4, 1.9, 1.3, 1.9])
    para(d,
         "partial 的語意是：已有的判據若超標仍會發報，但「沒有告警」不等於"
         "「已確認平靜」——缺少的那個判據可能單獨觸發。", size=11, after=8)

    # ── 三、資料源 ──────────────────────────────────────────────
    heading(d, "三、資料源")
    table(d,
          ["來源", "產品", "本案用途"],
          [["TACC / UCAR CDAAC", "福衛七號 leoOrb 精密定軌（SP3-c，60 秒節奏）",
            "反演阻力衰減率與密度增強倍數"],
           ["TACC / UCAR CDAAC", "福衛七號 scn1c2 掩星閃爍指數",
            "GNSS 網域的 S4 實測判據"],
           ["中央氣象署 SWOO", "TWTEC、TWDI", "區域 TEC 與地磁擾動"],
           ["NOAA SWPC", "X 射線、質子、Kp 預報、D-RAP 等", "HF 網域判據與預報基準"],
           ["GFZ Potsdam", "Kp／ap、Hp30", "地磁指數"],
           ["CelesTrak", "太空天氣快照、地球定向參數（EOP）", "F10.7／Ap 驅動、框架轉換"]],
          widths=[1.5, 2.7, 2.1])
    para(d,
         "福衛七號資料的正式取用條款確認，屬合作夥伴（成功大學團隊）之工作項目，"
         "預計於兩年期計畫內完成。目前所依據的 Release Memorandum "
         "（F7C2_SpWx_DataRelease_5.pdf）為技術性釋出說明，未見註冊要求或"
         "再散布限制；引用時須標註 TACC 與 UCAR CDAAC。", size=11, after=8)

    # ── 四、為何非用精密定軌不可 ────────────────────────────────
    heading(d, "四、為何非用精密定軌不可：TLE 走不通")
    para(d,
         "同樣的量測先以 TLE 做過，失敗。本專案既有的 drag_residual() 以"
         "中位數回歸自校準等效彈道係數（B_eff = median(−Δa／s)），實跑庫內"
         "六顆福衛七號 2024-01-01 至 2026-08-16 的 TLE：")
    table(d,
          ["項目", "結果"],
          [["每日阻力衰減預期", "約 40 公尺"],
           ["TLE 半長軸逐日殘差 RMS", "約 450–680 公尺"],
           ["訊噪比", "約 0.07"],
           ["六顆同型同軌道衛星的 B_eff 差異", "30%（真實差異不可能這麼大）"],
           ["B_eff 逐年變化", "2024→2026 接近翻倍，方向與推進劑消耗應造成的相反"]],
          widths=[3.2, 3.1])
    para(d,
         "最後一列是關鍵：B_eff 若隨推進劑消耗變化，質量減少應使其下降，"
         "實測方向相反。這代表 B_eff 吸收的是經驗模式隨太陽活動漂移的偏差。"
         "把它再拿去反演密度，得到的是模式尺度的密度——循環論證。")
    para(d,
         "leoOrb 的弧段重疊一致性實測為 3D RMS 0.25 公尺（徑向 0.15、沿跡 0.17、"
         "法向 0.08），平滑後的估計雜訊約 5 公尺。問題不在方法，在輸入資料的精度。",
         after=8)

    # ── 五、方法 ────────────────────────────────────────────────
    heading(d, "五、方法")
    heading(d, "5.1　彈道係數為何會消掉", 2)
    para(d,
         "近圓軌道的長期衰減為 da/dt = −B · ρ · sqrt(μa)，其中 B = Cd·A/m。"
         "福衛七號的投影面積、阻力係數與乾重皆未公開（且太陽能板為可調傾角，"
         "僅憑姿態四元數也算不出投影面積），故無法解出絕對密度。"
         "但取同一顆衛星兩個時段的比值時 B 完全消掉：")
    para(d, "ρ_now ／ ρ_quiet ＝ (da/dt)_now ／ (da/dt)_quiet",
         size=12, bold=True, align=A.CENTER, after=6)
    para(d,
         "分子分母皆為觀測量，因此不需要任何非公開的衛星參數，也不繼承經驗模式"
         "在暴時的偏差。這是本參數能宣稱 inference = observed 的依據。", after=8)

    heading(d, "5.2　條件一：平均窗必須是軌道週期的整數倍", 2)
    para(d,
         "密切半長軸的短週期振幅達 1.7 公里，遠大於每日數十公尺的衰減。"
         "若以 6 小時或日曆日取平均，短週期會混疊成假的長期趨勢。"
         "本案以軌道週期（約 96.31 分鐘）為窗做中心移動平均，"
         "實測可把逐點變異由 1118 公尺壓到 13 公尺，抑制約 85 倍。", after=8)

    heading(d, "5.3　條件二：基線必須先扣掉太陽通量", 2)
    para(d,
         "衰減率是絕對速率，基線隨太陽週期變動極大：實測 2024-05 平靜期約 "
         "59 m/日，2026-08 僅約 7.7 m/日，相差近八倍。而測試窗（2024-04-29 至 "
         "05-19）內 F10.7 由 132 漲到 238，衰減率隨之由 40 漲到 55 m/日。")
    para(d,
         "初版以滾動分位數當基線，結果平靜期的「增強倍數」是 1.4–1.65 而非 1.0，"
         "L1 持續誤觸發——太陽驅動被算成了地磁事件效應。改為以地磁寧靜的分箱"
         "（Kp < 4）擬合 log(衰減率) ～ a + b·F10.7，再以該式預測「同一 F10.7、"
         "地磁寧靜」應有的衰減率，觀測值除以它即為地磁造成的增強。")
    para(d,
         "擬合結果 b ≈ 0.0058，即 F10.7 每增加 100 sfu 密度約 ×1.78，物理合理。"
         "衰減率、F10.7、Kp 皆為觀測量，故此步驟不引入模式。", after=8)

    heading(d, "5.4　跨衛星彙整與機動處理", 2)
    para(d,
         "五至六顆同型、同軌道面的衛星在同一時刻承受同一組大氣，真實密度變化"
         "必為共模；單顆的機動則是離群值。因此**取值用中位數（對離群穩健）、"
         "偵測用標準差（對離群敏感）**——兩者的要求相反。離散度超過 0.25 時"
         "標 suspect 而非丟棄。", after=8)

    # ── 六、結果 ────────────────────────────────────────────────
    heading(d, "六、結果")
    heading(d, "6.1　2024-05 Gannon 事件", 2)
    table(d,
          ["時間（UT）", "衰減率 m/日", "增強倍數", "備註"],
          [["5/8–5/10 12:00", "約 59", "0.91–1.13", "暴前基線"],
           ["5/10 18:00", "93", "1.33", "SSC 約 17:05"],
           ["5/11 00:00", "215", "3.30", ""],
           ["5/11 06:00", "266", "4.01", "峰值"],
           ["5/11 12:00", "224", "3.53", ""],
           ["5/11 18:00", "155", "2.34", ""],
           ["5/12 06:00", "74", "1.07", "恢復"]],
          widths=[1.6, 1.5, 1.4, 1.8])
    para(d,
         "六顆衛星在峰值的離散度僅約 10%，非雜訊。相對照的是：同一窗的平靜週"
         "（2026-08）五顆殘差相關係數接近零、90% 變異為個別雜訊——"
         "平靜期看不到共模訊號是預期的，不能據以否定方法。", size=11, after=8)

    heading(d, "6.2　與經驗模式的比較：暴時響應被壓縮", 2)
    para(d,
         "以 MSIS 2.1 沿同一軌跡計算「同一 F10.7、地磁寧靜」為基準的 storm_ratio，"
         "與觀測增強倍數相除：")
    table(d,
          ["觀測增強倍數", "樣本數", "觀測中位", "模式中位", "觀測／模式"],
          [["< 1.5", "72", "1.01", "1.09", "0.85"],
           ["1.5 – 2.5", "8", "1.80", "1.21", "1.40"],
           ["2.5 – 3.5", "1", "3.30", "2.52", "1.31"],
           ["≥ 3.5", "2", "3.77", "2.30", "1.70"]],
          widths=[1.5, 1.0, 1.2, 1.2, 1.4])
    para(d,
         "即模式的暴時響應被壓縮：平靜時略微高估，實際增強愈大就愈低估。")
    para(d,
         "**能宣稱的是趨勢，不是絕對量。** 觀測側的 F10.7 迴歸基線與模式側的 "
         "storm_ratio 之間可能存在常數偏移，該偏移會平移整欄比值。"
         "但比值隨增強倍數單調上升這件事對常數偏移免疫，故「響應被壓縮」成立，"
         "而「平靜時高估 15%」不可單獨引用。", after=8)

    heading(d, "6.3　密度響應對 Kp 高度非線性", 2)
    table(d,
          ["Kp 區間", "樣本數", "增強倍數中位", "最大"],
          [["< 4（寧靜）", "61", "1.01", "1.98"],
           ["4 – 5", "7", "1.10", "1.97"],
           ["5 – 6（G1–G2）", "2", "1.21", "1.50"],
           ["6 – 7（G2–G3）", "6", "1.19", "1.49"],
           ["≥ 7（G3 以上）", "7", "2.34", "4.01"]],
          widths=[1.7, 1.2, 1.7, 1.2])
    para(d,
         "Kp 6–7 幾乎不推升 550 公里的密度。同一 21 天窗內，Kp 組規則觸發 15 次、"
         "實測組僅 4 次且全落在 Gannon——**Kp 代理在中等擾動時明顯過度告警**。"
         "這是實測判據存在的直接理由，也是本次工作最可能引起討論的一點。", after=8)

    heading(d, "6.4　門檻標定", 2)
    para(d,
         "以該窗 83 個 6 小時分箱中的 65 個地磁寧靜箱標定：中位 1.01、P90 1.42、"
         "P95 1.60、P98 1.77、最大 1.98。")
    table(d,
          ["等級", "門檻", "駐留", "遲滯解除", "Gannon 是否觸發"],
          [["L1", "≥ 2.0", "連續 2 箱", "< 1.6", "是（連續 3 箱）"],
           ["L2", "≥ 2.5", "—", "< 2.0", "是"],
           ["L3", "≥ 3.2", "—", "< 2.6", "是"],
           ["L4", "≥ 4.0", "—", "< 3.0", "是（峰值 4.01）"]],
          widths=[0.9, 1.1, 1.3, 1.3, 1.7])
    para(d,
         "L1 取 2.0 恰在寧靜期實測最大值之上，該窗內單箱與連續兩箱的誤報皆為 0。"
         "另加駐留條件是因為雜訊不相關而暴期會持續，這比單純調高門檻更能保住"
         "靈敏度。calibrated 維持 false。", size=11, after=8)

    heading(d, "6.5　附帶成果：TLE 誤差評估", 2)
    para(d,
         "同一份精密定軌可作為 TLE 誤差的真值參考。以 5 顆衛星、7 天、"
         "9,923 個比對點測得（550 km、傾角 24 度）：徑向 RMS 0.142 km、"
         "沿跡 0.731 km、法向 0.170 km。誤差以沿跡為主且僅沿跡明顯隨 TLE 齡期"
         "成長（0–3 h 為 0.43 km，24–48 h 為 1.46 km）。此結果已另撰文件提供"
         "Sat_TraingDataExtension 專案，用於取代該案目前假設的協方差。", after=8)

    # ── 七、已修正的錯誤 ────────────────────────────────────────
    heading(d, "七、本次工作中已發現並修正的錯誤")
    para(d,
         "列出這些是為了讓審閱者判斷結論的可信程度。三個都不是筆誤，"
         "而是會改變結論的實質錯誤。", after=4)
    table(d,
          ["曾經的錯誤", "如何發現", "修正"],
          [["以 6 小時重取樣算出衰減 39 m/日、訊噪比 150",
            "改以軌道週期為窗重算後得 7.7 m/日；與 da/dt = −B·ρ·v·a 的理論值"
            "約 6 m/日 相符",
            "確立平均窗須為軌道週期整數倍；理論值成為驗收依據"],
           ["判定某顆衛星「七天內無機動跡象」（已寫入交付文件）",
            "同樣是混疊所致。正確平滑後該星半長軸在 6 小時內上升 24 公尺，"
            "其餘四顆持續衰減",
            "更正交付文件；該次機動也解釋了該星沿跡 TLE 誤差為其餘四顆兩倍"],
           ["以中位數絕對偏差（MAD）當跨衛星離散度指標",
            "自撰測試失敗。n 只有五、六顆時，除非過半衛星偏離否則 MAD 恆為零",
            "改為取值用中位數、偵測用標準差；正好偵測不到單顆機動是最需要"
            "偵測的情形"]],
          widths=[2.0, 2.4, 1.9])
    para(d,
         "此外，「以滾動分位數當基線」的錯誤（第 5.3 節）之所以被抓到，"
         "是因為本專案 orbit_drag.density_ratio 的註解早已載明該陷阱："
         "「基準二：同一 F10.7、地磁寧靜……否則會把太陽週期當成事件效應」。",
         size=11, after=8)

    # ── 八、待確認事項 ──────────────────────────────────────────
    heading(d, "八、待確認事項（請同儕優先給意見）")
    bullets(d, [
        "**門檻標定的樣本太薄**。僅一個 21 天窗、65 個地磁寧靜樣本，"
        "且寧靜期最大值 1.98 幾乎貼著 L1 門檻 2.0，樣本增加後極可能超過。"
        "leoOrb 涵蓋 2019 年起，建議回填多個暴期把樣本做厚後再定版。",
        "**單一高度、單一傾角**。550 公里、傾角 24 度。極軌與太陽同步軌道的"
        "密度響應是否可沿用同一組門檻，需要另行驗證。",
        "**太陽能板可調傾角對 B 的影響未量化**。取比值時假設 B 在比較窗內固定；"
        "板角隨 β 角變化的週期是數月，故短窗（數日）內應可接受，但未實測。",
        "**觀測與模式兩條基線之間的常數偏移未確定**。目前只敢宣稱單調趨勢。"
        "要能宣稱絕對偏差，需要獨立的絕對密度校準來源。",
        "**尾隨特性的作業意涵**。衰減率標在 t 的值描述 t−6h 到 t，"
        "不適合當暴起始指標。與 Kp 組並用是否為最佳配置，值得討論。",
        "**是否要把 RHO_RATIO（觀測／模式）列為正式交付參數**。"
        "目前由分析工具算出、未入庫。若入庫則須決定基線窗的預設值。",
        "**福衛七號資料的正式取用條款**，屬合作夥伴工作項目，"
        "投稿前需確認引用方式。",
    ])

    # ── 九、相依性失效 ──────────────────────────────────────────
    heading(d, "九、附帶發現：一個不產生錯誤數字、只產生缺失數字的失效")
    para(d,
         "本次工作中，環境內的 astropy 7.1.1 與 numpy 2.5.0 不相容，"
         "其 time 與 coordinates 模組完全無法載入。兩個互相獨立的斷點："
         "numpy 2.0 移除的陣列函式，以及 astropy 跨套件邊界引用的私有符號"
         "（numpy 2.5 已改名）。常見的相容墊片只解決第一個。")
    para(d,
         "更嚴重的是診斷方向被帶偏：某呼叫端攔到 ImportError 後重拋為"
         "「請安裝 astropy」，而該套件明明已安裝。實際後果是 TLE 對精密星曆的"
         "殘差比對管線無法執行、結果表持續空白且不報錯。")
    para(d,
         "升級至 astropy 8.0.1 後解決（numpy 未動，其他九個共用套件無回歸）。"
         "本專案另實作獨立的 TEME↔ITRF 轉換，以 2,636 個實際軌道點對釘 "
         "astropy 8.0.1，平均差 0.081 公尺、最大 0.275 公尺，並將參考值固化為"
         "不需連網亦不需該套件的回歸測試。")
    para(d,
         "**建議的一般結論**：相依性驗證應納入作業型太空天氣軟體的資料品質規範，"
         "因為這類失效不產生錯的數字，只產生缺的數字——而缺的數字在儀表板上"
         "與「平靜」難以區分。這正是本系統把「沒資料」與「沒事」分開的同一個理由。",
         after=8)

    # ── 十、重現 ────────────────────────────────────────────────
    heading(d, "十、重現指令")
    table(d,
          ["宣稱", "指令"],
          [["擷取 leoOrb 並產生兩個參數",
            "python -m services.ingest.run --source tacc_leoorb --date 2024.140"],
           ["觀測與模式的密度比較",
            "python tools/density_obs_vs_model.py --start 2024-04-29 --end 2024-05-20"],
           ["分級規則行為（含駐留與遲滯）", "python -m pytest tests/test_tacc_leoorb.py"],
           ["框架轉換對 astropy 的回歸釘樁", "python -m pytest tests/test_frames.py"],
           ["端到端鏈路與無前視偏差回放", "python tools/e2e_demo.py --as-of 2024-05-11T00:00Z"]],
          widths=[2.4, 4.0])
    para(d,
         "leoOrb 為手動或排程來源，排除於背景自動更新之外（單日 1.4–2.2 MB，"
         "且需連續多日方能算出衰減率）。", size=11)

    return d


def main() -> None:
    out = "docs/AOSWA2026_Abstract_SpaceWeather_App_中文detail.docx"
    build().save(out)
    print(f"已產生 {out}")


if __name__ == "__main__":
    main()
